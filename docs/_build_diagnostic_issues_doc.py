# Generates Word doc with diagnostic warnings, failures, and incomplete items (ASCII style).
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = Path(__file__).resolve().parent / "MiddinInnovatie-Diagnostic-Issues.docx"

MONO = "Consolas"
MONO_SIZE = Pt(9.5)

REPORT_LINES = [
    "+======================================================================+",
    "|     MIDDIN INNOVATIE - DIAGNOSTIC ISSUES (WARN / FAIL / INCOMPLETE)  |",
    "+======================================================================+",
    f"|  Date : {datetime.now().strftime('%Y-%m-%d %H:%M')}",
    f"|  Repo : {ROOT}",
    "|  Ver  : 0.9.6 (versionCode 15)",
    "+======================================================================+",
    "",
    "+----------------------------------------------------------------------+",
    "| OVERALL STATUS:  NEEDS ATTENTION  (1 fail, 2 warnings, 23 pass)     |",
    "+----------------------------------------------------------------------+",
    "",
    "+----------------------------------------------------------------------+",
    "| [!!] WARNINGS (2)                                                    |",
    "+----------------------------------------------------------------------+",
    "|                                                                      |",
    "|  1. RSS / FALLBACK URL CHECKS                                        |",
    "|     Status : 14 ok, 2 failed (of 16 URLs)                           |",
    "|     Impact : App uses static fallback for news; not a crash          |",
    "|                                                                      |",
    "|     Failed URLs:                                                     |",
    "|       [403] https://www.healthcarefinancenews.com/rss.xml            |",
    "|             External server returned: (403) Forbidden                |",
    "|                                                                      |",
    "|       [--]  https://www.rijksoverheid.nl/onderwerpen/e-health        |",
    "|             Connection closed unexpectedly                           |",
    "|                                                                      |",
    "|     Action : Optional - replace or remove blocked feeds in             |",
    "|              InnovationRssSources.kt if you want 16/16 live          |",
    "|                                                                      |",
    "|  2. LARGE UNCOMMITTED CHANGE SET (GIT)                               |",
    "|     Status : Many local changes not committed or pushed              |",
    "|     Includes:                                                        |",
    "|       - Changelog auto-update feature (new files, not on GitHub)     |",
    "|       - Gemini removed (Android + desktop)                           |",
    "|       - Bluetooth removed (desktop only)                             |",
    "|       - Modified release scripts, strings, MorePanel, etc.           |",
    "|                                                                      |",
    "|     Action : Review git status, commit, push when ready              |",
    "|                                                                      |",
    "+----------------------------------------------------------------------+",
    "",
    "+----------------------------------------------------------------------+",
    "| [XX] FAILURES (1)                                                    |",
    "+----------------------------------------------------------------------+",
    "|                                                                      |",
    "|  1. ANDROID GRADLE BUILD FAILED                                      |",
    "|     Task   : :app:compileDebugKotlin (kspDebugKotlin)                |",
    "|     Cause  : SSL / PKIX certificate error on this PC                 |",
    "|                                                                      |",
    "|     Error excerpt:                                                   |",
    "|       Could not resolve org.jetbrains.kotlin:kotlin-stdlib-jdk8:1.8.21",
    "|       Got SSL handshake exception during request                     |",
    "|       PKIX path building failed: unable to find valid certification  |",
    "|       path to requested target                                       |",
    "|                                                                      |",
    "|     Repos affected:                                                  |",
    "|       - https://dl.google.com/dl/android/maven2/                     |",
    "|       - https://repo.maven.apache.org/maven2/                        |",
    "|                                                                      |",
    "|     Verdict: Environment / network issue, NOT an app code bug        |",
    "|                                                                      |",
    "|     Actions to try:                                                  |",
    "|       [ ] Check corporate proxy or antivirus SSL inspection          |",
    "|       [ ] Import company CA into Java truststore (JDK used by Gradle)|",
    "|       [ ] Build APK on another machine or via Android Studio         |",
    "|       [ ] Run: .\\gradlew.bat :app:compileDebugKotlin --no-daemon    |",
    "|                                                                      |",
    "+----------------------------------------------------------------------+",
    "",
    "+----------------------------------------------------------------------+",
    "| [??] INCOMPLETE / MISSING (3)                                        |",
    "+----------------------------------------------------------------------+",
    "|                                                                      |",
    "|  1. CHANGELOG.JSON NOT ON GITHUB                                     |",
    "|     Local  : OK  releases/changelog.json                             |",
    "|              OK  app/src/main/assets/changelog.json                  |",
    "|              OK  desktop/src/main/resources/changelog.json           |",
    "|     Remote : HTTP 404 Not Found                                      |",
    "|     URL    : https://raw.githubusercontent.com/pgraafland356-debug/  |",
    "|              MiddinInnovatie/main/releases/changelog.json            |",
    "|                                                                      |",
    "|     Impact : Remote auto-changelog will not update until pushed      |",
    "|     Action : git add releases/changelog.json + sync targets          |",
    "|              git commit && git push origin main                      |",
    "|                                                                      |",
    "|  2. CHANGELOG FEATURE NOT FULLY ROLLED OUT                           |",
    "|     Scripts exist: update-changelog.ps1, sync-changelog.ps1         |",
    "|     Wired in     : publish-github-release.ps1                        |",
    "|     Missing step : Push changelog.json to GitHub (see item 1)        |",
    "|                                                                      |",
    "|  3. ANDROID APK CANNOT BE REBUILT ON THIS PC                         |",
    "|     Related to failure #1 (SSL). Desktop builds work offline.        |",
    "|     Installed app is still 0.9.6 / code 15 (matches project).        |",
    "|                                                                      |",
    "+----------------------------------------------------------------------+",
    "",
    "+----------------------------------------------------------------------+",
    "| QUICK FIX CHECKLIST                                                  |",
    "+----------------------------------------------------------------------+",
    "|                                                                      |",
    "|  Priority 1 - Push changelog to GitHub:                              |",
    "|    .\\scripts\\sync-changelog.ps1                                    |",
    "|    git add releases/changelog.json app/src/main/assets/ ...          |",
    "|    git commit -m \"Add auto-updating changelog feed\"                |",
    "|    git push origin main                                              |",
    "|                                                                      |",
    "|  Priority 2 - Commit pending feature work:                           |",
    "|    git status                                                        |",
    "|    (review Gemini/Bluetooth removal + changelog changes)             |",
    "|                                                                      |",
    "|  Priority 3 - Fix SSL for Android builds (this machine):             |",
    "|    See failure #1 actions above                                      |",
    "|                                                                      |",
    "|  Priority 4 - RSS feeds (optional):                                  |",
    "|    Update or remove 2 failing news source URLs                       |",
    "|                                                                      |",
    "|  Re-run full diagnostic:                                             |",
    "|    .\\scripts\\run-full-diagnostic.ps1                               |",
    "|                                                                      |",
    "+----------------------------------------------------------------------+",
    "",
    "+======================================================================+",
    "|  SUMMARY TABLE                                                       |",
    "+======================================================================+",
    "|  Category      | Count | Severity                                    |",
    "+----------------+-------+---------------------------------------------+",
    "|  Warnings      |   2   | Low-Medium (RSS + uncommitted git)          |",
    "|  Failures      |   1   | High (Android build blocked on this PC)    |",
    "|  Incomplete    |   3   | Medium (changelog not live on GitHub)       |",
    "+----------------+-------+---------------------------------------------+",
    "|  Desktop app   |  OK   | Installed, compiles, updater present        |",
    "|  Release feed  |  OK   | latest.json live, SHA-256 matches           |",
    "+======================================================================+",
]


def set_mono(run, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = MONO
    run._element.rPr.rFonts.set(qn("w:eastAsia"), MONO)
    run.font.size = MONO_SIZE
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_ascii_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        color = None
        if line.startswith("| [!!]") or "WARNINGS" in line:
            color = RGBColor(0xCC, 0x88, 0x00)
        elif line.startswith("| [XX]") or "FAILURES" in line:
            color = RGBColor(0xCC, 0x00, 0x00)
        elif line.startswith("| [??]") or "INCOMPLETE" in line:
            color = RGBColor(0x00, 0x66, 0xCC)
        elif "NEEDS ATTENTION" in line:
            color = RGBColor(0xCC, 0x00, 0x00)
        set_mono(run, color=color)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Middin Innovatie - Diagnostic Issues Report")
    tr.font.size = Pt(14)
    tr.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("ASCII-style export of warnings, failures, and incomplete items")
    sr.font.size = Pt(10)
    sr.font.italic = True

    doc.add_paragraph()
    add_ascii_block(doc, REPORT_LINES)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    nr = note.add_run(
        "Generated by docs/_build_diagnostic_issues_doc.py from run-full-diagnostic.ps1 results."
    )
    nr.font.size = Pt(9)
    nr.font.italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
