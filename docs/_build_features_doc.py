# One-off script: generates Dutch Word doc with feature/fix ideas. Run: python _build_features_doc.py
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "Functies-en-verbeteringen-Middin-Innovatie.docx"


def main() -> None:
    doc = Document()
    t = doc.add_heading("Middin Innovatie-app — ideeën voor functies en fixes", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(
        "Dit document is een inventarisatie op basis van de huidige codebase (Compose, Room, Ktor, RSS-nieuws, "
        "lokale chat, producten met camera, instellingen). Gebruik het als backlog of discussiedocument; "
        "prioriteiten zijn niet vastgelegd."
    )

    doc.add_heading("Mogelijke nieuwe functies", level=1)

    features = [
        "Pushmeldingen voor nieuwe innovatienieuws of RSS-updates (nu vooral een testnotificatie).",
        "Offline modus voor het nieuwsoverzicht: laatst opgehaalde artikelen tonen zonder netwerk.",
        "Zoeken en filteren binnen collectief geheugen en chatgeschiedenis.",
        "Delen van producten of nieuwsartikelen via het Android-deelblad.",
        "Biometrische login (vingerafdruk / gezicht) na eerste serveraanmelding.",
        "Wachtwoord tonen/verbergen en optie “ingelogd blijven” op het inlogscherm.",
        "Meertalige changelog en updates: Nederlandse teksten naast of in plaats van Engels in de UI.",
        "Onboarding na eerste login (korte uitleg: tabs, geheugen, producten, API-instellingen).",
        "Widget voor startscherm met laatste nieuws of snelkoppeling naar chat.",
        "Export van collectief geheugen (bijv. tekst of JSON) voor archief of rapportage.",
        "Koppeling met echte backend-chat (RemoteChatRepository) als optionele modus naast lokale product-chat.",
        "Beheerderscherm: gebruikers of sessies (alleen zinvol als de API dat ondersteunt).",
        "Donkere modus verfijnen (contrast, kleur van links) specifiek voor het nieuwsoverzicht en kaarten.",
        "Toegankelijkheid: grotere tekst, TalkBack-labels en focusvolgorde nalopen op alle schermen.",
        "Play-internal test / automatische versiecontrole via een eenvoudige “nieuwe versie beschikbaar”-URL.",
    ]
    for line in features:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Mogelijke fixes en technische verbeteringen", level=1)

    fixes = [
        "Netwerk (fysiek toestel): cleartext HTTP is nu beperkt tot 10.0.2.2, localhost en 127.0.0.1. "
        "Voor een telefoon op Wi-Fi naar een pc op het LAN is vaak een extra domain-config of HTTPS op de dev-server nodig.",
        "Release-build: standaard API-URL is een placeholder (api.example.com); controleren vóór publicatie en ProGuard-regels voor serialisatie.",
        "Beveiliging: lokale dev-accounts (hardcoded whitelist) uitsluitend in debug; documenteren voor het team en nooit in release inschakelen.",
        "API-pad en loginveld (username vs e-mail) staan in Gradle; foutmeldingen kunnen duidelijker als de server 404/401 geeft.",
        "Gradle: bij dex-merge OutOfMemoryError helpt vaak minder parallelisme of meer geheugen; eventueel CI-documentatie bijwerken.",
        "Gemini-sleutel in DataStore: overwegen of gevoelige sleutels beter in EncryptedSharedPreferences / Android Keystore horen.",
        "Bluetooth-scherm: controleren of permissies en UX op Android 12+ volledig zijn afgestemd op scan/connect-flows.",
        "Welkomstscherm: optie “Overslaan” of opnieuw tonen via instellingen (nu alleen via app-gegevens wissen).",
        "E2E- of screenshot-tests voor kritieke paden (inloggen, home, product toevoegen) uitbreiden naast bestaande unit-tests.",
        "Consistente merkkleuren over alle schermen (bijv. secundaire accenten voor chips en links waar nu nog standaard Material-kleuren staan).",
    ]
    for line in fixes:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Al aanwezig (kort ter referentie)", level=1)
    existing = [
        "Tabs: Home, Collectief geheugen, Chat, Producten, Meer.",
        "Home: innovatienieuws (RSS), filters per bron, top producten.",
        "Producten: catalogus, camera/ML Kit, lokaal opslaan.",
        "Inloggen: server (Ktor) of lokaal (debug); API-basis-URL instelbaar; thema licht/donker/systeem.",
        "Meer: Instellingen, Changelog, Updates, Info, Over ons, Credits, Gemini-assistent, Bluetooth.",
        "Sessie: uitloggen, wissen bij scherm uit en bij swipe uit recente apps (service).",
        "Eenmalig welkomstscherm vóór inloggen (merk).",
    ]
    for line in existing:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    run = foot.add_run(f"Bestand gegenereerd voor projectmap docs/. Zie codebase onder app/src/main/.")
    run.italic = True
    run.font.size = Pt(9)

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
